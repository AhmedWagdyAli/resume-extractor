from docx import Document
from docx.shared import Inches
import os


def add_image_watermark(doc_path, image_path, output_path):
    # Check if the input file exists
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Input file not found: {doc_path}")

    # Check if the image file exists
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Image file not found: {image_path}")

    # Load the Word document
    doc = Document(doc_path)

    # Access the document's first section
    section = doc.sections[0]

    # Modify the header for the watermark
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]

    # Add an image to the header (acts as a watermark)
    run = header_paragraph.add_run()
    run.add_picture(image_path, width=Inches(6))  # Adjust size as needed

    # Save the modified document
    doc.save(output_path)


# Usage
input_doc_path = "./template.docx"  # Update this with the correct file path
output_doc_path = "./output_with_image.docx"
image_path = "./watermark.png"

add_image_watermark(input_doc_path, image_path, output_doc_path)
