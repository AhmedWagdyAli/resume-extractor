import os
from docx import Document
import logging
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

# Configure logging
logging.basicConfig(level=logging.DEBUG)


class CVProcessor:
    def __init__(self):
        self.doc = Document()

    def add_section(self, doc, title, items, bullet_points=False):
        doc.add_heading(title, level=1)
        if isinstance(items, list):
            for item in items:
                if isinstance(item, dict):
                    # Add each key-value pair from the dictionary as a paragraph
                    for key, value in item.items():
                        doc.add_paragraph(f"{key}: {value}")
                    doc.add_paragraph()  # Add a blank line between items for clarity
                else:
                    doc.add_paragraph(
                        item, style="List Bullet" if bullet_points else None
                    )
        elif isinstance(items, dict):
            for key, value in items.items():
                doc.add_paragraph(f"{key}: {value}")
        else:
            doc.add_paragraph(items)

    def add_personal_info(self, doc, data, coded):
        if coded == "code":
            personal_info = [
                f"Name: {'**'}",
                f"Email: {'**@gmail.com'}",
                f"Phone 1: {'****'}",
                f"Phone 2: {'Not provided'}",
                f"Address: {data['address'] or 'Not provided'}",
                f"City: {data['city'] or 'Not provided'}",
                f"LinkedIn: {'**@linkedin.com' or 'Not provided'}",
                f"Professional Experience in Years: {data['professional_experience_in_years']}",
                f"Highest Education: {data['highest_education']}",
            ]
        elif coded == "name":
            personal_info = [
                f"Name: {data['name']}",
                f"Email: {'**@gmail.com'}",
                f"Phone 1: {'****'}",
                f"Phone 2: {'Not provided'}",
                f"Address: {data['address'] or 'Not provided'}",
                f"City: {data['city'] or 'Not provided'}",
                f"LinkedIn: {'**@linkedin.com' or 'Not provided'}",
                f"Professional Experience in Years: {data['professional_experience_in_years']}",
                f"Highest Education: {data['highest_education']}",
            ]
        else:
            personal_info = [
                f"Name: {data['name']}",
                f"Email: {data['email']}",
                f"Phone 1: {data['phone_1']}",
                f"Phone 2: {data['phone_2'] or 'Not provided'}",
                f"Address: {data['address'] or 'Not provided'}",
                f"City: {data['city'] or 'Not provided'}",
                f"LinkedIn: {data['linkedin'] or 'Not provided'}",
                f"Professional Experience in Years: {data['professional_experience_in_years']}",
                f"Highest Education: {data['highest_education']}",
            ]
        for info in personal_info:
            doc.add_paragraph(info)

    def format_experience(self, experience_list, doc):
        if not isinstance(experience_list, list) or not experience_list:
            doc.add_paragraph("Not provided")
            return

        for item in experience_list:
            exp_paragraph = doc.add_paragraph()
            for key, value in item.items():
                if value:
                    # Bold organisation_name; normal for others.
                    if key == "organisation_name":
                        run = exp_paragraph.add_run(
                            f"{key.replace('_', ' ').title()}: {value}\n"
                        )
                        run.bold = True
                    else:
                        exp_paragraph.add_run(
                            f"{key.replace('_', ' ').title()}: {value}\n"
                        )
            exp_paragraph.add_run("\n")  # Extra newline between experiences

    def format_projects(self, proj_list, doc):
        if not isinstance(proj_list, list) or not proj_list:
            doc.add_paragraph("Not provided")
            return

        for project in proj_list:
            proj_paragraph = doc.add_paragraph()
            for key, value in project.items():
                if value:
                    # Bold the title field
                    if key == "title":
                        run = proj_paragraph.add_run(
                            f"{key.replace('_', ' ').title()}: {value}\n"
                        )
                        run.bold = True
                    else:
                        proj_paragraph.add_run(
                            f"{key.replace('_', ' ').title()}: {value}\n"
                        )
            proj_paragraph.add_run("\n")  # Extra newline between projects

    def format_skills(self, skills_text, doc):
        if not skills_text or not isinstance(skills_text, list):
            doc.add_paragraph("N/A")
            return
        for skill in skills_text:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(skill.strip())

    def format_education(self, education_list, doc):
        if not isinstance(education_list, list) or not education_list:
            doc.add_paragraph("N/A")
            return
        for item in education_list:
            edu_paragraph = doc.add_paragraph()
            for key, value in item.items():
                if value:
                    edu_paragraph.add_run(f"{key.replace('_', ' ').title()}: {value}\n")
            edu_paragraph.add_run("\n")

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """
        Inserts a new paragraph after the given paragraph and returns it.
        """
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style:
            new_para.style = style
        return new_para

    def replace_placeholders(self, doc, data, template_type):
        """Replace placeholders in a Word document and insert special content under the correct headers."""

        def add_personal_info(data, template_type):
            if template_type == "code":
                return "\n".join(
                    [
                        f"Name: {'**'}",
                        f"Email: {'**@gmail.com'}",
                        f"Phone 1: {'****'}",
                        f"Phone 2: {'Not provided'}",
                        f"Address: {data.get('address', 'Not provided')}",
                        f"City: {data.get('city', 'Not provided')}",
                        f"LinkedIn: {'**@linkedin.com'}",
                        f"Professional Experience in Years: {data.get('professional_experience_in_years', 'Not provided')}",
                        f"Highest Education: {data.get('highest_education', 'Not provided')}",
                    ]
                )
            elif template_type == "name":
                return "\n".join(
                    [
                        f"Name: {data.get('name', 'Not provided')}",
                        f"Email: {'**@gmail.com'}",
                        f"Phone 1: {'****'}",
                        f"Phone 2: {'Not provided'}",
                        f"Address: {data.get('address', 'Not provided')}",
                        f"City: {data.get('city', 'Not provided')}",
                        f"LinkedIn: {'**@linkedin.com'}",
                        f"Professional Experience in Years: {data.get('professional_experience_in_years', 'Not provided')}",
                        f"Highest Education: {data.get('highest_education', 'Not provided')}",
                    ]
                )
            else:
                return "\n".join(
                    [
                        f"Name: {data.get('name', 'Not provided')}",
                        f"Email: {data.get('email', 'Not provided')}",
                        f"Phone 1: {data.get('phone_1', 'Not provided')}",
                        f"Phone 2: {data.get('phone_2', 'Not provided')}",
                        f"Address: {data.get('address', 'Not provided')}",
                        f"City: {data.get('city', 'Not provided')}",
                        f"LinkedIn: {data.get('linkedin', 'Not provided')}",
                        f"Professional Experience in Years: {data.get('professional_experience_in_years', 'Not provided')}",
                        f"Highest Education: {data.get('highest_education', 'Not provided')}",
                    ]
                )

        def insert_experience(after_paragraph, exp_list):
            """Insert formatted professional experience paragraphs after the given paragraph."""
            for item in exp_list:
                new_p = self.insert_paragraph_after(after_paragraph)  # Updated call
                for key, value in item.items():
                    if value:
                        if key == "organisation_name":
                            run = new_p.add_run(
                                f"{key.replace('_', ' ').title()}: {value}\n"
                            )
                            run.bold = True
                        else:
                            new_p.add_run(f"{key.replace('_', ' ').title()}: {value}\n")

        def insert_projects(after_paragraph, proj_list):
            """Insert formatted project paragraphs after the given paragraph."""
            for project in proj_list:
                new_p = self.insert_paragraph_after(after_paragraph)  # Updated call
                for key, value in project.items():
                    if value:
                        if key == "title":
                            run = new_p.add_run(
                                f"{key.replace('_', ' ').title()}: {value}\n"
                            )
                            run.bold = True
                        else:
                            new_p.add_run(f"{key.replace('_', ' ').title()}: {value}\n")

        def insert_skills(after_paragraph, skills_list):
            """Insert each skill as a bullet point after the given paragraph."""
            for skill in skills_list:
                new_p = self.insert_paragraph_after(
                    after_paragraph, style="List Bullet"
                )  # Updated call
                new_p.add_run(skill.strip())

        def insert_education(after_paragraph, education_list):
            """Insert formatted education paragraphs after the given paragraph."""
            for item in education_list:
                new_p = self.insert_paragraph_after(after_paragraph)  # Updated call
                for key, value in item.items():
                    if value:
                        new_p.add_run(f"{key.replace('_', ' ').title()}: {value}\n")

        def convert_value(key, value):
            """Convert individual key-value pairs for normal placeholders."""
            if key == "certifications" and isinstance(value, list):
                return "\n".join(cert["name"] for cert in value)
            if key in ["name", "email", "phone_1", "phone_2"]:
                if template_type == "code":
                    if key == "name":
                        return "*****"
                    elif key == "email":
                        return "****@gmail.com"
                    elif key in ["phone_1", "phone_2"]:
                        return "****" if key == "phone_1" else "Not provided"
                elif template_type == "name":
                    if key == "name":
                        return data.get("name", "Not provided")
                    elif key == "email":
                        return "**@gmail.com"
                    elif key in ["phone_1", "phone_2"]:
                        return "****" if key == "phone_1" else "Not provided"
            return str(value) if value is not None else "Not provided"

        # Get a copy of the original paragraphs to avoid processing newly added ones.
        original_paragraphs = list(doc.paragraphs)

        # Process special placeholders by inserting content immediately after the paragraph that contains the placeholder.
        for paragraph in original_paragraphs:
            if "{{professional_experience}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{professional_experience}}", ""
                )
                insert_experience(paragraph, data.get("professional_experience", []))
            if "{{projects}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{projects}}", "")
                insert_projects(paragraph, data.get("projects", []))
            if "{{skills}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{skills}}", "")
                insert_skills(paragraph, data.get("skills", []))
            if "{{education}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{education}}", "")
                insert_education(paragraph, data.get("education", []))

        # Process remaining placeholders in the original paragraphs.
        for paragraph in original_paragraphs:
            for key, value in data.items():
                if key in [
                    "professional_experience",
                    "projects",
                    "skills",
                    "education",
                ]:
                    continue  # already handled
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        placeholder, convert_value(key, value)
                    )

        # Process placeholders in tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(
                                placeholder, convert_value(key, value)
                            )

        # Add combined personal info
        combined_personal_info = add_personal_info(data, template_type)
        for paragraph in original_paragraphs:
            if "{{personal_info}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{personal_info}}", combined_personal_info
                )

    def fill_cv_template(self, data, output_path, template_type):
        try:
            project_root = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(project_root, "template.docx")
            doc = Document(template_path)
            self.replace_placeholders(doc, data, template_type)
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)
            doc.save(output_path)
            logging.debug(f"Document created: {output_path}")
        except Exception as e:
            logging.error(f"Error filling the template: {e}")

    def fill_template(self, data, output_path):
        self.fill_cv_template(data, output_path, template_type="normal")

    def fill_coded_template(self, data, output_path):
        self.fill_cv_template(data, output_path, template_type="code")

    def fill_name_template(self, data, output_path):
        self.fill_cv_template(data, output_path, template_type="name")
