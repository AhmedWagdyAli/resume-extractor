from docx import Document
import random
import string


def add_section(doc, heading, items, bullet_points=False):
    doc.add_heading(heading, level=1)
    if bullet_points:
        for item in items:
            doc.add_paragraph(item, style="ListBullet")
    else:
        for item in items:
            doc.add_paragraph(item)


def add_personal_info(doc, data, coded):
    if coded == "coded" | coded == "name":
        personal_info = [
            f"Name: {''.join(random.choices(string.ascii_uppercase + string.digits, k=10))}",
            f"Email: {'**@gmail.com'}",
            f"Phone 1: {'****'}",
            f"Address: {data['address'] or 'Not provided'}",
            f"City: {data['city'] or 'Not provided'}",
            f"LinkedIn: {'**@linkedin.com' or 'Not provided'}",
            f"Professional Experience in Years: {data['professional_experience_in_years']}",
            f"Highest Education: {data['highest_education']}",
        ]
    else:
        personal_info = [
            f"Name: {data['name']}",
            f"Email: {data['email']}",
            f"Phone 1: {data['phone_1']}",
            f"Phone 2: {data['phone_2'] or 'Not provided'}",
            f"Address: {data['address'] or 'Not provided'}",
            f"City: {data['city'] or 'Not provided'}",
            f"LinkedIn: {data['linkedin'] or 'Not provided'}",
            f"Professional Experience in Years: {data['professional_experience_in_years']}",
            f"Highest Education: {data['highest_education']}",
        ]
    for info in personal_info:
        doc.add_paragraph(info)


def fill_template_common(data, output_path, coded=False):
    try:
        doc = Document()
        add_personal_info(doc, data, coded=coded)
        add_section(doc, "Skills", data["skills"], bullet_points=True)
        add_section(doc, "Education", data["education"])
        add_section(doc, "Professional Experience", data["professional_experience"])
        doc.save(output_path)
        print(f"Document created: {output_path}")
    except Exception as e:
        print(f"Error filling the template: {e}")


def fill_template(data, output_path):
    fill_template_common(data, output_path, coded="normal")


def fill_coded_template(data, output_path):
    fill_template_common(data, output_path, coded="coded")


def fill_name_template(data, output_path):
    fill_template_common(data, output_path, coded="name")
